import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent


def set_default_font(document: Document, font_name: str = "Times New Roman", font_size_pt: int = 10) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str, size: int = 12) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    p.space_before = Pt(8)
    p.space_after = Pt(4)


def add_name_block(document: Document, content: dict) -> None:
    name = content.get("name", "Name")
    preferred = content.get("preferred_name")
    display_name = f"{name} ({preferred})" if preferred else name
    subtitle = f"{content.get('title', '')} · {content.get('affiliation', '')}".strip(" ·")

    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(display_name)
    r1.bold = True
    r1.font.size = Pt(18)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)

    contact = " · ".join(
        [
            x
            for x in [
                content.get("email"),
                content.get("phone"),
                f"WeChat: {content.get('wechat_id')}" if content.get("wechat_id") else "",
            ]
            if x
        ]
    )
    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(contact).font.size = Pt(10)

    if content.get("website"):
        p4 = document.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run(f"· Homepage: {content['website']}").font.size = Pt(10)


def add_bullets(document: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        document.add_paragraph(item, style=style)


def add_first_page_gap(document: Document, pt: float = 4) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def parse_start_year(period: str) -> int:
    if not period:
        return -1
    chunk = period.split("-")[0].strip()
    digits = "".join(ch for ch in chunk if ch.isdigit())
    if len(digits) >= 4:
        return int(digits[:4])
    return -1


def build_docx(content: dict, out_path: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_name_block(doc, content)

    add_heading(doc, "Professional Summary")
    doc.add_paragraph(content.get("summary", ""))
    add_first_page_gap(doc, 6)

    add_heading(doc, "Core Strengths")
    add_bullets(doc, content.get("strengths", []))
    add_first_page_gap(doc, 6)

    add_heading(doc, "Education")
    add_bullets(doc, content.get("education", []))
    add_first_page_gap(doc, 6)

    add_heading(doc, "Real-World AI Impact")
    for item in content.get("research_impact", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item.get('theme', '')}. ").bold = True
        p.add_run(f"Focus: {item.get('focus', '')} Contribution: {item.get('contribution', '')}")
    add_first_page_gap(doc, 6)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Selected Publications (Condensed)")
    add_bullets(doc, content.get("selected_publications", []))

    add_heading(doc, "Experience")
    experience = sorted(content.get("experience", []), key=lambda x: parse_start_year(x.get("period", "")), reverse=True)
    for exp in experience:
        doc.add_paragraph(f"{exp.get('period', '')}  {exp.get('title', '')}, {exp.get('organization', '')}")
        add_bullets(doc, exp.get("bullets", []), style="List Bullet 2")

    add_heading(doc, "Teaching & Mentoring")
    add_bullets(doc, content.get("teaching", []))

    add_heading(doc, "Professional Service")
    add_bullets(doc, content.get("service", []))

    add_heading(doc, "Awards")
    add_bullets(doc, content.get("awards", []))

    add_heading(doc, "Technical & Collaboration Skills")
    add_bullets(doc, content.get("skills", []))

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    content_path = ROOT / "cv_content_consulting.json"
    out_docx = ROOT / "research_cv_targeted_consulting.docx"
    content = json.loads(content_path.read_text(encoding="utf-8"))
    build_docx(content, out_docx)
    print(f"Wrote {out_docx}")


if __name__ == "__main__":
    main()
