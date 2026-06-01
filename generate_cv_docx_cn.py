import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent


def _set_default_font(document: Document, font_name: str = "Times New Roman", font_size_pt: int = 10) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.space_before = Pt(10)
    p.space_after = Pt(4)


def _add_subheading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    p.space_before = Pt(6)
    p.space_after = Pt(2)


def _add_section_hint(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    p.space_before = Pt(0)
    p.space_after = Pt(4)


def _add_name_block(
    document: Document, name: str, subtitle: str, contact_line: str, homepage_line: str
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(18)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(contact_line)
    r3.font.size = Pt(10)

    if homepage_line:
        p4 = document.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(homepage_line)
        r4.font.size = Pt(10)


def _add_bullets(document: Document, items: list[str], style: str = "List Bullet") -> None:
    for it in items:
        document.add_paragraph(it, style=style)


def _fmt_month(month: str | None) -> str:
    if not month:
        return ""
    m = month.replace(".", "-")
    parts = m.split("-")
    if len(parts) >= 2:
        return f"{parts[0]}.{parts[1]}"
    return month


def _edu_line_zh(edu: dict) -> str:
    start = _fmt_month(edu.get("start"))
    end = edu.get("end", "")
    end = "至今" if str(end).lower() == "present" else _fmt_month(str(end))
    date = f"{start}–{end}".strip("–")
    degree_raw = str(edu.get("degree", "")).strip()
    degree = degree_raw
    major_map = {
        "Computer Science and Engineering": "计算机科学与工程",
        "Computer Science": "计算机科学与工程",
        "Big Data Technology": "大数据技术",
        "Statistics": "统计学",
    }
    if "Ph.D. in" in degree_raw:
        major = degree_raw.replace("Ph.D. in", "").strip()
        degree = f"博士（{major_map.get(major, major)}）"
    elif "M.Sc. in" in degree_raw:
        major = degree_raw.replace("M.Sc. in", "").strip()
        degree = f"硕士（{major_map.get(major, major)}）"
    elif "B.S. in" in degree_raw:
        major = degree_raw.replace("B.S. in", "").strip()
        degree = f"学士（{major_map.get(major, major)}）"
    inst = str(edu.get("institution", "")).replace(
        "Hong Kong University of Science and Technology", "香港科技大学"
    ).replace("Beijing Normal-Hong Kong Baptist University (BNBU)", "北师香港浸会大学")
    return f"{date}  {degree}，{inst}".strip()


def _work_citation(work: dict) -> str:
    authors = ", ".join(work.get("authors", []))
    title = work.get("title", "")
    venue = work.get("venue", "")
    year = work.get("year", "")
    venue_text = "投稿中" if venue == "In submission" else venue
    suffix = f"{venue_text}, {year}" if venue_text else str(year)
    base = f"{authors}. {title}. {suffix}."
    notes = work.get("notes") or []
    if notes:
        translated_notes = [
            "* Equal contribution (alphabetical order)" if "Equal contribution (alphabetical order)" in n else n
            for n in notes
        ]
        base += " " + " ".join(translated_notes)
    return base.replace("  ", " ").strip()


def _translate_degree_notes(note: str) -> str:
    translated = (
        note.replace("Huawei PhD Fellowship (recipient)", "华为博士奖学金获得者")
        .replace("CGPA:", "绩点：")
        .replace("Rank:", "排名：")
        .replace("Graduate First-Class Honour, Hong Kong Baptist University", "香港浸会大学一级荣誉毕业")
    )
    return translated.replace("： ", "：")


def _translate_summary(summary: str) -> str:
    mapping = {
        "Multimodal LLMs may rely on language priors rather than pertinent visual evidence, especially on long documents. I explore agentic perception frameworks that gather evidence iteratively to improve accuracy and efficiency.": "多模态大语言模型（Multimodal LLMs）在长文档场景中可能过度依赖语言先验，而忽略关键视觉证据。我关注 agentic perception 框架，通过迭代式证据采集来同时提升准确率与效率。",
        "Controllable editing requires precise spatial and semantic guidance without costly retraining. I develop training-free methods that combine structural control with flexible prompt guidance.": "可控图像编辑需要在不进行高成本重训练的前提下，实现精确的空间与语义引导。我研究 training-free 方法，将结构化控制与灵活的 prompt guidance 相结合。",
        "Foundation models can lose robustness during fine-tuning and fail under distribution shift. I study causal mechanisms behind domain generalization and design training objectives that anchor decisions to invariant, generalizable features.": "基础模型（Foundation models）在微调过程中可能出现鲁棒性下降，并在分布偏移下失效。我研究 domain generalization 背后的因果机制，并设计训练目标，使模型决策锚定在稳定且可泛化的特征上。",
        "Deep classifiers often rely on spurious correlations rather than causally relevant visual evidence. My work develops explanation methods that diagnose these misaligned dependencies and surface discriminative, human-interpretable rationales.": "深度分类器常依赖虚假相关而非因果相关的视觉证据。我研究 explainable AI 方法，用于诊断这类偏置依赖，并提供具有区分性且可解释的决策依据。",
    }
    return mapping.get(summary, summary)


def _translate_description(description: str) -> str:
    mapping = {
        "Replaces fixed-resolution, single-pass pipelines with iterative perception that selectively acquires high-resolution crops on demand, advancing the accuracy–efficiency Pareto frontier.": "以迭代式感知替代固定分辨率、单次前向流程，按需获取高分辨率局部区域，从而推进准确率与效率的 Pareto 前沿。",
        "Enables selective edge-based structural control with dual-prompt guidance for training-free, controllable image editing.": "通过 selective edge-based structural control 与 dual-prompt guidance，实现 training-free 的可控图像编辑。",
        "Logit Attribution Matching (LAM) anchors predictions to domain-invariant causal features by matching logit attributions across semantic-sharing pairs.": "通过在语义相近样本对之间匹配 logit attribution，Logit Attribution Matching (LAM) 将模型预测锚定到 domain-invariant 的因果特征。",
        "Combats robustness vanishing during foundation-model fine-tuning by jointly optimizing empirical risk with worst-case risk estimated via CLIP and LLM-generated visual descriptions.": "通过联合优化经验风险与最坏情形风险（基于 CLIP 与 LLM 生成的视觉描述估计），缓解 foundation model 在微调阶段的鲁棒性退化问题。",
        "Estimates the causal effect of semantic patches on Vision Transformer predictions, moving beyond correlational saliency maps.": "估计语义图块对 Vision Transformer 预测结果的因果效应，超越传统相关性 saliency map 的解释范式。",
        "Introduces Contrastive Whole-Output Explanation (CWOX), which explains a model's top-K labels by systematically contrasting visually confusable competitors.": "提出 Contrastive Whole-Output Explanation (CWOX)，通过系统性对比视觉上易混淆的类别，解释模型 top-K 输出。",
        "Proposes a diagnostic measure for assessing how well a model captures the structure of training examples.": "提出用于评估模型是否有效捕捉训练样本结构的诊断指标。",
    }
    return mapping.get(description, description)


def build_docx(content: dict, out_path: Path) -> None:
    doc = Document()
    _set_default_font(doc, "Times New Roman", 10)

    name = content.get("name", "Name")
    preferred = content.get("preferred_name")
    display_name = f"{name} ({preferred})" if preferred else name

    title_zh = "计算机科学博士生"
    aff_zh = "香港科技大学（HKUST）"
    subtitle = f"{title_zh} · {aff_zh}"
    website = content.get("website")
    homepage_line = f"· 个人主页：{website}" if website else ""
    wechat = content.get("wechat_id")
    wechat_text = f"微信：{wechat}" if wechat else ""
    contact_line = " · ".join(
        [
            x
            for x in [
                content.get("email"),
                content.get("phone"),
                wechat_text,
            ]
            if x
        ]
    )
    _add_name_block(doc, display_name, subtitle, contact_line, homepage_line)

    _add_heading(doc, "经历")
    for exp in content.get("experience", []):
        start = str(exp.get("start", "")).strip()
        end = str(exp.get("end", "")).strip()
        date = f"{start}–{end}".strip("–")
        role = str(exp.get("role", "")).replace("Research Intern", "研究实习生")
        org = exp.get("organization", "").strip().replace(
            "Huawei Hong Kong AI Framework & Data Technologies Lab", "华为香港AI框架与数据技术实验室"
        )
        line = f"{date}  {role}，{org}".strip("，")
        doc.add_paragraph(line)
    for edu in content.get("education", []):
        doc.add_paragraph(_edu_line_zh(edu))
        for note in edu.get("notes", []) or []:
            doc.add_paragraph(_translate_degree_notes(note), style="List Bullet 2")
        advisor = edu.get("advisor")
        if advisor:
            doc.add_paragraph(f"导师：{advisor}", style="List Bullet 2")

    statement = content.get("research_statement")
    if statement:
        _add_heading(doc, "研究陈述")
        doc.add_paragraph(
            "我的研究聚焦于深度视觉与视觉-语言模型在真实场景中的应用，重点关注可解释性、泛化能力、"
            "多模态大语言模型（MLLMs）的计算效率，以及图像编辑中的可控性。"
            "我致力于构建诊断工具来分析模型当前依赖的关键信号，并设计有针对性的机制，"
            "引导模型学习更具因果相关性、更可信且更高效的行为。"
        )

    _add_heading(doc, "Research Interests & Selected Work")
    _add_section_hint(doc, "Organized by research theme, each with representative publications and brief contributions.")
    for idx, area in enumerate(content.get("research_areas", []), start=1):
        _add_subheading(doc, f"主题 {idx}：{area.get('title', '')}")
        summary = area.get("summary")
        if summary:
            p_summary = doc.add_paragraph()
            p_summary.add_run("Focus: ").bold = True
            p_summary.add_run(summary)
        doc.add_paragraph("Representative work:", style="List Bullet")
        for work in area.get("works", []):
            p = doc.add_paragraph(style="List Bullet 2")
            p.add_run(_work_citation(work)).bold = True
            description = work.get("description")
            if description:
                doc.add_paragraph(description, style="List Bullet 3")

    teaching = content.get("teaching", [])
    if teaching:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _add_heading(doc, "教学经历")
        for entry in teaching:
            role = str(entry.get("role", "")).replace("Teaching Assistant", "助教")
            terms = entry.get("terms", [])
            doc.add_paragraph(f"{role}：{', '.join(terms)}")

    _add_heading(doc, "学术服务")
    for block in content.get("service", []):
        role = block.get("role")
        if role:
            doc.add_paragraph(str(role).replace("Conference reviewer / PC member", "会议审稿人 / 程序委员会成员"))
        _add_bullets(doc, block.get("items", []))

    _add_heading(doc, "荣誉与奖项")
    awards = [
        x.replace("Huawei PhD Fellowship (HKUST)", "华为博士奖学金（HKUST）")
        .replace("MSc Big Data Technology Top Students Award (HKUST), 2020", "大数据技术硕士优秀学生奖（HKUST），2020")
        .replace("School of Engineering Excellent Student Scholarship (HKUST), 2020", "工学院优秀学生奖学金（HKUST），2020")
        .replace("College First Class Scholarship (BNBU), 2016–2019", "学院一等奖学金（BNBU），2016–2019")
        for x in content.get("awards", [])
    ]
    _add_bullets(doc, awards)

    _add_heading(doc, "技术技能")
    skills = content.get("skills", {}) or {}
    skill_lines: list[str] = []
    if skills.get("ml_ai"):
        skill_lines.append("ML & AI: " + ", ".join(skills["ml_ai"]))
    if skills.get("programming"):
        skill_lines.append("编程与框架: " + ", ".join(skills["programming"]))
    if skills.get("tools"):
        skill_lines.append("工具与平台: " + ", ".join(skills["tools"]))
    for ln in skill_lines:
        doc.add_paragraph(ln, style="List Bullet")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    content_path = ROOT / "cv_content.json"
    out_docx = ROOT / "research_cv_updated_cn.docx"
    content = json.loads(content_path.read_text(encoding="utf-8"))
    build_docx(content, out_docx)
    print(f"Wrote {out_docx}")


if __name__ == "__main__":
    main()
